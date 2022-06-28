from contextlib import nullcontext
import io
from urllib import response
from .serializers import CoachSerializer, NewOrganiserUserSerializer, NewCoachUserSerializer, OrganiserSerializer, UserSerializer, EventSerializer
from django.http import StreamingHttpResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import Coach, Event, Organiser, User
from rest_framework.authtoken.models import Token
from datetime import datetime, timedelta
from django.core.mail import send_mail
from docx import Document

approved_emails = ['vladcoroian2001@gmail.com', 'og519@ic.ac.uk', 'adiboroica235@gmail.com']

class EventView(APIView):
    def get(self, request, format=None):
        events = Event.objects.filter(
            organiser_user=request.user).order_by('date')
        serializer = EventSerializer(events, many=True)
        return Response(serializer.data)

    def post(self, request):
        request.data['organiser_user_id'] = request.user.pk
        serializer = EventSerializer(data=request.data)
        if serializer.is_valid(raise_exception=ValueError):
            event = serializer.create(validated_data=request.data)
            organiser = event.organiser_user
            request.data['pk'] = event.pk
            request.data.pop('organiser_user_id')
            for coach in request.user.organiser.favourites.all():
                if coach.email in approved_emails:
                    send_mail(
                        'New Job Offer',
                        'An organiser wants you to take a look at a potential opportunity. \n' +
                        f'{organiser.first_name} {organiser.last_name} would like to invite you to ' +
                        f'apply for {event.name}, on {event.date}, at {event.location}.\n' + 
                        'To get more information or to apply for this opportunity open KeepPlaying.' +
                        '\n\n Best, \n Keep Playing Team',
                        'drp@keep_playing.com',
                        [coach.email],
                        fail_silently=False,
                    )
            return Response(
                request.data,
                status=status.HTTP_201_CREATED
            )
        return Response(
            {
                "error": True,
                "error_msg": serializer.error_messages,
            },
            status=status.HTTP_400_BAD_REQUEST
        )

    def delete(self, request, pk, format=None):
        try:
            event = Event.objects.get(pk=pk)
            event.delete()
        except Event.DoesNotExist:
            return Response(
                {
                    "error": True,
                    "error_msg": "Event does not exist",
                },
                status=status.HTTP_400_BAD_REQUEST
            )
        return Response({'message': 'Deleted'})

    def patch(self, request, pk, format=None):
        try:
            event = Event.objects.get(pk=pk)
            serializer = EventSerializer(
                event, data=request.data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return Response(
                    serializer.data,
                    status=status.HTTP_202_ACCEPTED
                )
        except Event.DoesNotExist:
            return Response(
                {
                    "error": True,
                    "error_msg": "Event does not exist",
                },
                status=status.HTTP_400_BAD_REQUEST
            )


class AcceptOfferView(APIView):
    def patch(self, request, pk, coach_pk, format=None):
        try:
            event = Event.objects.get(pk=pk)
            serializer = EventSerializer(
                event, data=request.data, partial=True)
            event.coach_user = User.objects.get(pk=coach_pk)
            if event.coach_user.email in approved_emails:
                send_mail(
                    'Your offer has been accepted!',
                    f'You have been accepted for {event.name}, on {event.date}, at {event.location}. Open Keep Playing for more details.\n' +
                    '\n\n Best, \n Keep Playing Team',
                    'drp@keep_playing.com',
                    [event.coach_user.email],
                    fail_silently=False,
                )
            if serializer.is_valid():
                serializer.save()
                return Response(
                    serializer.data,
                    status=status.HTTP_202_ACCEPTED
                )
        except Event.DoesNotExist:
            return Response(
                {
                    "error": True,
                    "error_msg": "Event does not exist",
                },
                status=status.HTTP_400_BAD_REQUEST
            )


class CoachCancelEventView(APIView):
    def patch(self, request, pk, format=None):
        try:
            event = Event.objects.get(pk=pk)
            serializer = EventSerializer(
                event, data=request.data, partial=True)
            event.coach = False
            event.coach_user = None
            event.offers.remove(self.request.user)
            if event.organiser_user.email in approved_emails:
                send_mail(
                    f'A {event.role} has cancelled!',
                    f'The {event.role} for {event.name}, on {event.date} has cancelled. ' + 
                    'Don\'t worry! We have already triggered another search. ' + 
                    'Open Keep Playing for more details.' +
                    '\n\n Best, \n Keep Playing Team',
                    'drp@keep_playing.com',
                    [event.organiser_user.email],
                    fail_silently=False,
                )
            if serializer.is_valid():
                serializer.save()
                return Response(
                    serializer.data,
                    status=status.HTTP_202_ACCEPTED
                )
        except Event.DoesNotExist:
            return Response(
                {
                    "error": True,
                    "error_msg": "Event does not exist",
                },
                status=status.HTTP_400_BAD_REQUEST
            )

class CoachUnapplyView(APIView):
    def patch(self, request, pk, format=None):
        try:
            event = Event.objects.get(pk=pk)
            serializer = EventSerializer(
                event, data=request.data, partial=True)
            event.offers.remove(request.user)
            if serializer.is_valid():
                serializer.save()
                return Response(
                    serializer.data,
                    status=status.HTTP_202_ACCEPTED
                )
        except Event.DoesNotExist:
            return Response(
                {
                    "error": True,
                    "error_msg": "Event does not exist",
                },
                status=status.HTTP_400_BAD_REQUEST
            )


class CoachOrganiserView(APIView):
    def get(self, request, pk, format=None):
        user = User.objects.get(pk=pk)
        serializer = UserSerializer(user, many=False)
        return Response(
            serializer.data,
            status=status.HTTP_202_ACCEPTED
        )

class CoachEventView(APIView):
    def get(self, request, pk, format=None):
        user = User.objects.get(pk=pk)
        serializer = UserSerializer(user, many=False)
        return Response(
            serializer.data,
            status=status.HTTP_202_ACCEPTED
        )

    def patch(self, request, pk, format=None):
        try:
            event = Event.objects.get(pk=pk)
            serializer = EventSerializer(
                event, data=request.data, partial=True)
            event.offers.add(request.user)
            if event.organiser_user.email in approved_emails:
                send_mail(
                    f'Offer received for {event.name}, on {event.date}',
                    f'You received a new offer from {request.user.first_name} {request.user.last_name}. ' +
                    'To get more information or to accept this offer open KeepPlaying.' +
                    '\n\n Best, \n Keep Playing Team',
                    'drp@keep_playing.com',
                    [event.organiser_user.email],
                    fail_silently=False,
                )
            if serializer.is_valid():
                serializer.save()
                return Response(
                    serializer.data,
                    status=status.HTTP_202_ACCEPTED
                )
        except Event.DoesNotExist:
            return Response(
                {
                    "error": True,
                    "error_msg": "Event does not exist",
                },
                status=status.HTTP_400_BAD_REQUEST
            )


class HelloView(APIView):
    def get(self, request, format=None):
        return Response("Hello {0}!".format(request.user))


class UsersRecordView(APIView):
    def get(self, format=None):
        users = User.objects.all()
        serializer = UserSerializer(users, many=True)
        return Response(serializer.data)


class UserRecordView(APIView):
    def get(self, format=None):
        user = self.request.user
        serializer = UserSerializer(user, many=False)
        response = serializer.data
        response['pk'] = user.pk
        return Response(response)

    def post(self, request):
        serializer = UserSerializer(data=request.data)
        if serializer.is_valid(raise_exception=ValueError):
            serializer.create(validated_data=request.data)
            return Response(
                serializer.data,
                status=status.HTTP_201_CREATED
            )
        return Response(
            {
                "error": True,
                "error_msg": serializer.error_messages,
            },
            status=status.HTTP_400_BAD_REQUEST
        )


class OrganiserView(APIView):
    def get(self, format=None):
        user = self.request.user
        serializer = OrganiserSerializer(user.organiser, many=False)
        return Response(serializer.data)

    def patch(self, request, format=None):
        organiser = self.request.user.organiser
        serializer = OrganiserSerializer(
            organiser, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(
                serializer.data,
                status=status.HTTP_202_ACCEPTED
            )
        return Response(
            {
                "error": True,
                "error_msg": "Organiser does not exist",
            },
            status=status.HTTP_400_BAD_REQUEST
        )


class OrganiserBlockCoachView(APIView):
    def patch(self, request, coach_pk, format=None):
        organiser = self.request.user.organiser
        organiser.blocked.add(coach_pk)
        serializer = OrganiserSerializer(organiser, many=False)
        return Response(
            serializer.data,
            status=status.HTTP_202_ACCEPTED
        )


class OrganiserUnblockCoachView(APIView):
    def patch(self, request, coach_pk, format=None):
        organiser = self.request.user.organiser
        organiser.blocked.remove(coach_pk)
        serializer = OrganiserSerializer(organiser, many=False)
        return Response(
            serializer.data,
            status=status.HTTP_202_ACCEPTED
        )


class OrganiserAddFavouriteCoachView(APIView):
    def patch(self, request, coach_pk, format=None):
        organiser = self.request.user.organiser
        organiser.favourites.add(coach_pk)
        serializer = OrganiserSerializer(organiser, many=False)
        return Response(
            serializer.data,
            status=status.HTTP_202_ACCEPTED
        )


class OrganiserRemoveFavouriteCoachView(APIView):
    def patch(self, request, coach_pk, format=None):
        organiser = self.request.user.organiser
        organiser.favourites.remove(coach_pk)
        serializer = OrganiserSerializer(organiser, many=False)
        return Response(
            serializer.data,
            status=status.HTTP_202_ACCEPTED
        )


class OrganiserEventsView(APIView):
    def get(self, request, format=None):
        events = Event.objects.filter(
            organiser_user=request.user).order_by('date')
        serializer = EventSerializer(events, many=True)
        return Response(serializer.data)


class CoachFeedView(APIView):
    def get(self, request, format=None):
        now = datetime.now()
        all_events = Event.objects.filter(date__gte=now, coach_user=None).order_by('date')
        valid_events = []
        for e in all_events:
            if not e.organiser_user.organiser.blocked.contains(self.request.user):
                valid_events.append(e)
        serializer = EventSerializer(valid_events, many=True)
        return Response(serializer.data)


class CoachUpcomingJobsView(APIView):
    def get(self, request, format=None):
        now = datetime.now()
        events = Event.objects.filter(coach_user=request.user).filter(
            date__gte=now).order_by('date')
        serializer = EventSerializer(events, many=True)
        return Response(serializer.data)
      
class VoteCoachView(APIView):
    def patch(self, request, event_pk, format=None):
        event = Event.objects.get(pk=event_pk)
        coach = event.coach_user.coach
        if not event.voted:
            event.voted = True
            event.save()
            coach.votes += 1
            coach.experience += request.data["experience"]
            coach.flexibility += request.data["flexibility"]
            coach.reliability += request.data["reliability"]
            coach.save()
        serializer = CoachSerializer(coach, many=False)
        return Response(serializer.data)

class CoachModelView(APIView):
    def get(self, request, coach_pk, format=None):
        user = User.objects.get(pk=coach_pk)
        coach = user.coach
        serializer = CoachSerializer(coach, many=False)
        return Response(serializer.data)

class ExportDocx(APIView):
    def build_document(self, event):
        document = Document() 

        # add a header
        document.add_heading(f"INVOICE FOR {event.name}, on {event.date}\n")

        # add a paragraph
        document.add_paragraph("This is a normal style paragraph")

        # add a paragraph within an italic text then go on with a break.
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.italic = True
        run.add_text("text will have italic style")
        run.add_break()
        
        return document

    def get(self, request, event_pk, *args, **kwargs):
        # create an empty document object
        event = Event.objects.get(pk=event_pk)
        document = self.build_document(event)
        # save document info
        buffer = io.BytesIO()
        document.save(buffer)  # save your memory stream
        buffer.seek(0)  # rewind the stream

        # put them to streaming content response 
        # within docx content_type
        response = StreamingHttpResponse(
            streaming_content=buffer,  # use the stream's content
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        response['Content-Disposition'] = 'attachment;filename=Invoice.docx'
        response["Content-Encoding"] = 'UTF-8'

        return response


class CreateCoachUser(APIView):
    def post(self, request, format=None):
        serializer = NewCoachUserSerializer(data=request.data)
        if serializer.is_valid(raise_exception=ValueError):
            serializer.create(validated_data=request.data)
            return Response({"message": "test"})
        return Response(
            {
                "error": True,
                "error_msg": serializer.error_messages,
            },
            status=status.HTTP_400_BAD_REQUEST
        )

class CreateOrganiserUser(APIView):
    def post(self, request, format=None):
        serializer = NewOrganiserUserSerializer(data=request.data)
        if serializer.is_valid(raise_exception=ValueError):
            serializer.create(validated_data=request.data)
            return Response({"message": "test"})
        return Response(
            {
                "error": True,
                "error_msg": serializer.error_messages,
            },
            status=status.HTTP_400_BAD_REQUEST
        )

class EventGetOrganiserView(APIView):
    def get(self, request, pk, format=None):
        event = Event.objects.get(pk=pk)
        user = event.organiser_user
        serializer = UserSerializer(user, many=False)
        return Response(serializer.data)